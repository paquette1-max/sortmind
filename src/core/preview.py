"""
File preview module for displaying file content before organizing.
Supports text files, images, PDFs, and documents.
"""
import logging
from pathlib import Path
from typing import Optional, Dict, Any, Union
from dataclasses import dataclass
from enum import Enum

logger = logging.getLogger(__name__)


class PreviewType(Enum):
    """Types of file previews supported."""
    TEXT = "text"
    IMAGE = "image"
    PDF = "pdf"
    DOCUMENT = "document"
    UNKNOWN = "unknown"


@dataclass
class FilePreview:
    """Preview data for a file."""
    file_path: Path
    preview_type: PreviewType
    preview_content: str  # Text content or description
    thumbnail_path: Optional[Path] = None  # For images/PDFs
    metadata: Dict[str, Any] = None
    error: Optional[str] = None


class PreviewManager:
    """Manages file preview generation."""
    
    # File type mappings
    TEXT_EXTENSIONS = {'.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', 
                       '.csv', '.log', '.ini', '.conf', '.sh', '.bash', '.zsh', '.yaml', 
                       '.yml', '.toml', '.cfg', '.rst', '.java', '.c', '.cpp', '.h', 
                       '.go', '.rs', '.rb', '.php', '.swift', '.kt', '.sql'}
    
    IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif', 
                        '.webp', '.heic', '.heif', '.raw', '.cr2', '.nef', '.arw'}
    
    PDF_EXTENSIONS = {'.pdf'}
    
    DOCUMENT_EXTENSIONS = {'.doc', '.docx', '.odt', '.rtf', '.pages', '.epub', '.mobi'}
    
    def __init__(self, max_text_preview: int = 500):
        """
        Initialize preview manager.
        
        Args:
            max_text_preview: Maximum characters to show for text previews
        """
        self.max_text_preview = max_text_preview
        self.logger = logger
    
    def get_preview(self, file_path: Union[str, Path]) -> FilePreview:
        """
        Get preview for a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            FilePreview object with preview data
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return FilePreview(
                file_path=file_path,
                preview_type=PreviewType.UNKNOWN,
                preview_content="",
                error=f"File not found: {file_path}"
            )
        
        try:
            preview_type = self._detect_preview_type(file_path)
            
            if preview_type == PreviewType.TEXT:
                return self._preview_text(file_path)
            elif preview_type == PreviewType.IMAGE:
                return self._preview_image(file_path)
            elif preview_type == PreviewType.PDF:
                return self._preview_pdf(file_path)
            elif preview_type == PreviewType.DOCUMENT:
                return self._preview_document(file_path)
            else:
                return self._preview_unknown(file_path)
                
        except Exception as e:
            self.logger.error(f"Failed to generate preview for {file_path}: {e}")
            return FilePreview(
                file_path=file_path,
                preview_type=PreviewType.UNKNOWN,
                preview_content="",
                error=str(e)
            )
    
    def _detect_preview_type(self, file_path: Path) -> PreviewType:
        """Detect the preview type based on file extension."""
        ext = file_path.suffix.lower()
        
        if ext in self.TEXT_EXTENSIONS:
            return PreviewType.TEXT
        elif ext in self.IMAGE_EXTENSIONS:
            return PreviewType.IMAGE
        elif ext in self.PDF_EXTENSIONS:
            return PreviewType.PDF
        elif ext in self.DOCUMENT_EXTENSIONS:
            return PreviewType.DOCUMENT
        else:
            # Try to detect text files by reading a sample
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    sample = f.read(1024)
                    # If we can read it as text and it looks like text
                    if sample and all(ord(c) < 128 or c.isprintable() for c in sample[:100]):
                        return PreviewType.TEXT
            except:
                pass
            return PreviewType.UNKNOWN
    
    def _preview_text(self, file_path: Path) -> FilePreview:
        """Generate preview for text files."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read(self.max_text_preview + 100)  # Read a bit extra
            
            # Truncate to max length
            if len(content) > self.max_text_preview:
                content = content[:self.max_text_preview] + "\n... [truncated]"
            
            # Get file stats
            stat = file_path.stat()
            
            metadata = {
                'size': stat.st_size,
                'size_human': self._format_size(stat.st_size),
                'line_count': content.count('\n') + 1,
                'char_count': len(content),
                'encoding': 'utf-8'  # We assume UTF-8 for preview
            }
            
            return FilePreview(
                file_path=file_path,
                preview_type=PreviewType.TEXT,
                preview_content=content,
                metadata=metadata
            )
            
        except Exception as e:
            return FilePreview(
                file_path=file_path,
                preview_type=PreviewType.TEXT,
                preview_content="",
                error=f"Failed to read text file: {e}"
            )
    
    def _preview_image(self, file_path: Path) -> FilePreview:
        """Generate preview for image files."""
        metadata = {'size': 0, 'size_human': '0 B'}
        exif_data = {}
        
        try:
            from PIL import Image
            from PIL.ExifTags import TAGS
            
            stat = file_path.stat()
            metadata['size'] = stat.st_size
            metadata['size_human'] = self._format_size(stat.st_size)
            
            with Image.open(file_path) as img:
                metadata['width'] = img.width
                metadata['height'] = img.height
                metadata['format'] = img.format
                metadata['mode'] = img.mode
                
                # Extract EXIF data
                try:
                    exif = img._getexif()
                    if exif:
                        for tag_id, value in exif.items():
                            tag = TAGS.get(tag_id, tag_id)
                            # Limit string length for display
                            if isinstance(value, str) and len(value) > 100:
                                value = value[:100] + "..."
                            exif_data[tag] = value
                        
                        # Extract key EXIF info
                        if 'DateTime' in exif_data:
                            metadata['date_taken'] = exif_data['DateTime']
                        if 'Make' in exif_data:
                            metadata['camera_make'] = exif_data['Make']
                        if 'Model' in exif_data:
                            metadata['camera_model'] = exif_data['Model']
                except:
                    pass
            
            # Build preview description
            preview_text = f"Image: {metadata['width']}x{metadata['height']} pixels\n"
            preview_text += f"Format: {metadata['format']}\n"
            preview_text += f"Mode: {metadata['mode']}\n"
            preview_text += f"Size: {metadata['size_human']}\n"
            
            if 'date_taken' in metadata:
                preview_text += f"Date Taken: {metadata['date_taken']}\n"
            if 'camera_make' in metadata:
                preview_text += f"Camera: {metadata['camera_make']}"
                if 'camera_model' in metadata:
                    preview_text += f" {metadata['camera_model']}"
                preview_text += "\n"
            
            metadata['exif'] = exif_data
            
            return FilePreview(
                file_path=file_path,
                preview_type=PreviewType.IMAGE,
                preview_content=preview_text,
                metadata=metadata
            )
            
        except ImportError:
            # PIL not available
            stat = file_path.stat()
            metadata['size'] = stat.st_size
            metadata['size_human'] = self._format_size(stat.st_size)
            
            return FilePreview(
                file_path=file_path,
                preview_type=PreviewType.IMAGE,
                preview_content=f"Image file ({metadata['size_human']})\n\nInstall Pillow for full preview support.",
                metadata=metadata,
                error="Pillow not installed"
            )
            
        except Exception as e:
            return FilePreview(
                file_path=file_path,
                preview_type=PreviewType.IMAGE,
                preview_content="",
                error=f"Failed to read image: {e}"
            )
    
    def _preview_pdf(self, file_path: Path) -> FilePreview:
        """Generate preview for PDF files."""
        metadata = {'size': 0, 'size_human': '0 B'}
        
        try:
            stat = file_path.stat()
            metadata['size'] = stat.st_size
            metadata['size_human'] = self._format_size(stat.st_size)
            
            # Try to use PyPDF2 or pdfplumber
            try:
                import pdfplumber
                
                with pdfplumber.open(file_path) as pdf:
                    metadata['page_count'] = len(pdf.pages)
                    
                    # Extract text from first page
                    if len(pdf.pages) > 0:
                        first_page = pdf.pages[0]
                        text = first_page.extract_text()
                        
                        if text:
                            # Truncate text
                            if len(text) > self.max_text_preview:
                                text = text[:self.max_text_preview] + "\n... [truncated]"
                            metadata['first_page_text'] = text
                        
                        # Get page dimensions
                        metadata['page_width'] = first_page.width
                        metadata['page_height'] = first_page.height
                
                preview_text = f"PDF Document\n"
                preview_text += f"Pages: {metadata['page_count']}\n"
                preview_text += f"Size: {metadata['size_human']}\n"
                
                if 'page_width' in metadata:
                    preview_text += f"Page Size: {metadata['page_width']:.0f} x {metadata['page_height']:.0f} pts\n"
                
                if 'first_page_text' in metadata:
                    preview_text += f"\n--- First Page Preview ---\n{metadata['first_page_text']}"
                
                return FilePreview(
                    file_path=file_path,
                    preview_type=PreviewType.PDF,
                    preview_content=preview_text,
                    metadata=metadata
                )
                
            except ImportError:
                pass
            
            # Fallback to PyPDF2
            try:
                from PyPDF2 import PdfReader
                
                reader = PdfReader(file_path)
                metadata['page_count'] = len(reader.pages)
                
                # Extract metadata
                if reader.metadata:
                    pdf_meta = reader.metadata
                    if pdf_meta.title:
                        metadata['title'] = pdf_meta.title
                    if pdf_meta.author:
                        metadata['author'] = pdf_meta.author
                    if pdf_meta.subject:
                        metadata['subject'] = pdf_meta.subject
                
                # Extract text from first page
                if len(reader.pages) > 0:
                    try:
                        text = reader.pages[0].extract_text()
                        if text and len(text) > self.max_text_preview:
                            text = text[:self.max_text_preview] + "\n... [truncated]"
                        metadata['first_page_text'] = text
                    except:
                        pass
                
                preview_text = f"PDF Document\n"
                preview_text += f"Pages: {metadata['page_count']}\n"
                preview_text += f"Size: {metadata['size_human']}\n"
                
                if 'title' in metadata:
                    preview_text += f"Title: {metadata['title']}\n"
                if 'author' in metadata:
                    preview_text += f"Author: {metadata['author']}\n"
                
                if 'first_page_text' in metadata and metadata['first_page_text']:
                    preview_text += f"\n--- First Page Preview ---\n{metadata['first_page_text']}"
                
                return FilePreview(
                    file_path=file_path,
                    preview_type=PreviewType.PDF,
                    preview_content=preview_text,
                    metadata=metadata
                )
                
            except ImportError:
                pass
            
            # No PDF library available
            return FilePreview(
                file_path=file_path,
                preview_type=PreviewType.PDF,
                preview_content=f"PDF file ({metadata['size_human']})\n\nInstall PyPDF2 or pdfplumber for full preview support.",
                metadata=metadata,
                error="PDF library not installed"
            )
            
        except Exception as e:
            return FilePreview(
                file_path=file_path,
                preview_type=PreviewType.PDF,
                preview_content="",
                error=f"Failed to read PDF: {e}"
            )
    
    def _preview_document(self, file_path: Path) -> FilePreview:
        """Generate preview for document files (docx, odt, etc)."""
        metadata = {'size': 0, 'size_human': '0 B'}
        
        try:
            stat = file_path.stat()
            metadata['size'] = stat.st_size
            metadata['size_human'] = self._format_size(stat.st_size)
            
            ext = file_path.suffix.lower()
            
            if ext == '.docx':
                try:
                    from docx import Document
                    
                    doc = Document(file_path)
                    
                    # Extract title if available
                    if doc.core_properties.title:
                        metadata['title'] = doc.core_properties.title
                    if doc.core_properties.author:
                        metadata['author'] = doc.core_properties.author
                    
                    # Extract first paragraph
                    paragraphs = []
                    for para in doc.paragraphs[:5]:  # First 5 paragraphs
                        if para.text.strip():
                            paragraphs.append(para.text.strip())
                    preview_text = f"Word Document (.docx)\n"
                    preview_text += f"Size: {metadata['size_human']}\n"
                    
                    if 'title' in metadata:
                        preview_text += f"Title: {metadata['title']}\n"
                    if 'author' in metadata:
                        preview_text += f"Author: {metadata['author']}\n"
                    
                    preview_text += f"Paragraphs: {len(doc.paragraphs)}\n\n"
                    
                    if paragraphs:
                        preview_text += "--- First Paragraphs ---\n"
                        preview_text += "\n\n".join(paragraphs[:3])  # First 3 non-empty paragraphs
                    
                    return FilePreview(
                        file_path=file_path,
                        preview_type=PreviewType.DOCUMENT,
                        preview_content=preview_text,
                        metadata=metadata
                    )
                    
                except ImportError:
                    pass
            
            # Generic document preview
            preview_text = f"Document ({ext})\n"
            preview_text += f"Size: {metadata['size_human']}\n"
            
            return FilePreview(
                file_path=file_path,
                preview_type=PreviewType.DOCUMENT,
                preview_content=preview_text,
                metadata=metadata
            )
            
        except Exception as e:
            return FilePreview(
                file_path=file_path,
                preview_type=PreviewType.DOCUMENT,
                preview_content="",
                error=f"Failed to read document: {e}"
            )
    
    def _preview_unknown(self, file_path: Path) -> FilePreview:
        """Generate preview for unknown file types."""
        try:
            stat = file_path.stat()
            metadata = {
                'size': stat.st_size,
                'size_human': self._format_size(stat.st_size),
                'modified': stat.st_mtime
            }
            
            preview_text = f"File type: {file_path.suffix or 'unknown'}\n"
            preview_text += f"Size: {metadata['size_human']}\n"
            
            return FilePreview(
                file_path=file_path,
                preview_type=PreviewType.UNKNOWN,
                preview_content=preview_text,
                metadata=metadata
            )
            
        except Exception as e:
            return FilePreview(
                file_path=file_path,
                preview_type=PreviewType.UNKNOWN,
                preview_content="",
                error=f"Failed to get file info: {e}"
            )
    
    def _format_size(self, size_bytes: int) -> str:
        """Format file size in human-readable format."""
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} PB"
