"""
Document content extraction using Strategy Pattern.
Each document type has its own extractor class for extensibility.
"""
import logging
from pathlib import Path
from typing import Optional, Dict, Any, Union, List, Protocol
from dataclasses import dataclass
from enum import Enum, auto

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Types of documents we can extract content from."""
    PDF = auto()
    IMAGE = auto()
    TEXT = auto()
    OFFICE = auto()
    UNKNOWN = auto()


@dataclass
class ExtractedContent:
    """Extracted content from a document."""
    file_path: Path
    doc_type: DocumentType
    text_content: str
    metadata: Dict[str, Any]
    extraction_method: str
    error: Optional[str] = None
    is_scanned_image: bool = False


class DocumentExtractor(Protocol):
    """Protocol for document extractors."""
    
    def can_extract(self, file_path: Path) -> bool:
        """Check if this extractor can handle the file."""
        ...
    
    def extract(self, file_path: Path, max_pages: int = 5, max_text_length: int = 8000) -> ExtractedContent:
        """Extract content from the document."""
        ...


class BaseExtractor:
    """Base class for extractors with common functionality."""
    
    @staticmethod
    def _format_size(size_bytes: int) -> str:
        """Format file size in human-readable format."""
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} PB"
    
    @staticmethod
    def _error_result(file_path: Path, doc_type: DocumentType, error: str) -> ExtractedContent:
        """Create an error result."""
        return ExtractedContent(
            file_path=file_path,
            doc_type=doc_type,
            text_content="",
            metadata={},
            extraction_method="none",
            error=error
        )


class PDFExtractor(BaseExtractor):
    """Extractor for PDF documents."""
    
    SUPPORTED_EXTENSIONS = {'.pdf'}
    
    def can_extract(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def extract(self, file_path: Path, max_pages: int = 5, max_text_length: int = 8000) -> ExtractedContent:
        """Extract text from PDF using available backends."""
        # Try PyMuPDF first
        result = self._try_pymupdf(file_path, max_pages, max_text_length)
        if result and not result.error:
            return result
        
        # Fallback to pdfplumber
        result = self._try_pdfplumber(file_path, max_pages, max_text_length)
        if result and not result.error:
            return result
        
        # Final fallback: OCR for scanned PDFs
        return self._try_ocr(file_path, max_pages, max_text_length)
    
    def _try_pymupdf(self, file_path: Path, max_pages: int, max_text_length: int) -> Optional[ExtractedContent]:
        """Try extracting with PyMuPDF."""
        try:
            import fitz
            
            text_parts = []
            metadata = {"pages": 0, "is_scanned": False}
            
            with fitz.open(file_path) as doc:
                metadata["pages"] = len(doc)
                
                for page_num, page in enumerate(doc[:max_pages]):
                    text = page.get_text()
                    if text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
                    else:
                        # Page appears to be scanned image
                        metadata["is_scanned"] = True
                        ocr_text = self._ocr_page(page, page_num)
                        if ocr_text:
                            text_parts.append(f"--- Page {page_num + 1} (OCR) ---\n{ocr_text}")
            
            full_text = "\n\n".join(text_parts)
            if len(full_text) > max_text_length:
                full_text = full_text[:max_text_length] + "\n... [truncated for length]"
            
            return ExtractedContent(
                file_path=file_path,
                doc_type=DocumentType.PDF,
                text_content=full_text,
                metadata=metadata,
                extraction_method="pymupdf" + ("+tesseract" if metadata["is_scanned"] else ""),
                is_scanned_image=metadata["is_scanned"]
            )
            
        except ImportError:
            return None
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}")
            return None
    
    def _try_pdfplumber(self, file_path: Path, max_pages: int, max_text_length: int) -> Optional[ExtractedContent]:
        """Try extracting with pdfplumber."""
        try:
            import pdfplumber
            
            text_parts = []
            metadata = {"pages": 0}
            
            with pdfplumber.open(file_path) as pdf:
                metadata["pages"] = len(pdf.pages)
                
                for i, page in enumerate(pdf.pages[:max_pages]):
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(f"--- Page {i + 1} ---\n{text}")
            
            full_text = "\n\n".join(text_parts)
            
            # If no text extracted, might be scanned PDF
            if not full_text.strip():
                return None  # Signal to try OCR
            
            if len(full_text) > max_text_length:
                full_text = full_text[:max_text_length] + "\n... [truncated for length]"
            
            return ExtractedContent(
                file_path=file_path,
                doc_type=DocumentType.PDF,
                text_content=full_text,
                metadata=metadata,
                extraction_method="pdfplumber"
            )
            
        except ImportError:
            return None
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")
            return None
    
    def _try_ocr(self, file_path: Path, max_pages: int, max_text_length: int) -> ExtractedContent:
        """Extract text from scanned PDF using OCR."""
        try:
            import pytesseract
            from pdf2image import convert_from_path
            from PIL import Image
            
            logger.info(f"Processing scanned PDF with OCR: {file_path.name}")
            
            images = convert_from_path(file_path, first_page=1, last_page=max_pages, dpi=200)
            
            text_parts = []
            for i, image in enumerate(images):
                text = pytesseract.image_to_string(image)
                if text.strip():
                    text_parts.append(f"--- Page {i + 1} ---\n{text}")
            
            full_text = "\n\n".join(text_parts)
            metadata = {"pages": len(images), "is_scanned": True}
            
            if len(full_text) > max_text_length:
                full_text = full_text[:max_text_length] + "\n... [truncated for length]"
            
            return ExtractedContent(
                file_path=file_path,
                doc_type=DocumentType.PDF,
                text_content=full_text,
                metadata=metadata,
                extraction_method="ocr",
                is_scanned_image=True
            )
            
        except ImportError:
            return self._error_result(
                file_path, DocumentType.PDF,
                "Scanned PDF OCR requires pdf2image and pytesseract. Install with: pip install pdf2image pytesseract"
            )
        except Exception as e:
            return self._error_result(file_path, DocumentType.PDF, f"OCR failed: {e}")
    
    def _ocr_page(self, page, page_num: int) -> Optional[str]:
        """OCR a single PDF page."""
        try:
            import fitz  # PyMuPDF - needed for Matrix
            import pytesseract
            from PIL import Image
            import io

            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            return pytesseract.image_to_string(img)
        except Exception:
            return None


class ImageExtractor(BaseExtractor):
    """Extractor for image files using OCR."""
    
    SUPPORTED_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif', '.webp'}
    
    def can_extract(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def extract(self, file_path: Path, max_pages: int = 5, max_text_length: int = 8000) -> ExtractedContent:
        """Extract text from image using OCR."""
        try:
            import pytesseract
            from PIL import Image
            
            with Image.open(file_path) as img:
                metadata = {
                    "width": img.width,
                    "height": img.height,
                    "mode": img.mode,
                    "format": img.format
                }
                
                text = pytesseract.image_to_string(img)
                
                return ExtractedContent(
                    file_path=file_path,
                    doc_type=DocumentType.IMAGE,
                    text_content=text,
                    metadata=metadata,
                    extraction_method="tesseract_ocr",
                    is_scanned_image=True
                )
                
        except ImportError:
            return self._error_result(
                file_path, DocumentType.IMAGE,
                "Image OCR requires pytesseract and Pillow. Install with: pip install pytesseract Pillow"
            )
        except Exception as e:
            return self._error_result(file_path, DocumentType.IMAGE, f"OCR failed: {e}")


class TextExtractor(BaseExtractor):
    """Extractor for text files."""
    
    SUPPORTED_EXTENSIONS = {'.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml',
                           '.csv', '.log', '.yaml', '.yml', '.toml', '.ini', '.cfg', '.rst',
                           '.java', '.c', '.cpp', '.h', '.go', '.rs', '.rb', '.php', '.swift'}
    
    def can_extract(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def extract(self, file_path: Path, max_pages: int = 5, max_text_length: int = 8000) -> ExtractedContent:
        """Extract content from text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read(max_text_length)
            
            if len(content) >= max_text_length:
                content = content[:max_text_length] + "\n... [truncated for length]"
            
            stat = file_path.stat()
            metadata = {
                "size": stat.st_size,
                "size_human": self._format_size(stat.st_size),
                "lines": content.count('\n') + 1
            }
            
            return ExtractedContent(
                file_path=file_path,
                doc_type=DocumentType.TEXT,
                text_content=content,
                metadata=metadata,
                extraction_method="direct_read"
            )
            
        except Exception as e:
            return self._error_result(file_path, DocumentType.TEXT, str(e))


class OfficeExtractor(BaseExtractor):
    """Extractor for Office documents."""
    
    SUPPORTED_EXTENSIONS = {'.docx', '.odt', '.doc', '.rtf'}
    
    def can_extract(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def extract(self, file_path: Path, max_pages: int = 5, max_text_length: int = 8000) -> ExtractedContent:
        """Extract text from Office documents."""
        ext = file_path.suffix.lower()
        
        if ext == '.docx':
            return self._extract_docx(file_path, max_text_length)
        
        # Generic fallback for other office formats
        stat = file_path.stat()
        metadata = {"size": stat.st_size, "size_human": self._format_size(stat.st_size)}
        
        return ExtractedContent(
            file_path=file_path,
            doc_type=DocumentType.OFFICE,
            text_content=f"Document type: {ext}\nSize: {metadata['size_human']}\n\nFull text extraction not available for this format.",
            metadata=metadata,
            extraction_method="basic"
        )
    
    def _extract_docx(self, file_path: Path, max_text_length: int) -> ExtractedContent:
        """Extract text from .docx files."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            content = "\n\n".join(paragraphs)
            if len(content) > max_text_length:
                content = content[:max_text_length] + "\n... [truncated]"
            
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }
            
            return ExtractedContent(
                file_path=file_path,
                doc_type=DocumentType.OFFICE,
                text_content=content,
                metadata=metadata,
                extraction_method="python-docx"
            )
            
        except ImportError:
            return self._error_result(
                file_path, DocumentType.OFFICE,
                ".docx extraction requires python-docx. Install with: pip install python-docx"
            )
        except Exception as e:
            return self._error_result(file_path, DocumentType.OFFICE, f"Failed to read document: {e}")


class GenericExtractor(BaseExtractor):
    """Fallback extractor for unknown file types."""
    
    def can_extract(self, file_path: Path) -> bool:
        """Always accepts as fallback."""
        return True
    
    def extract(self, file_path: Path, max_pages: int = 5, max_text_length: int = 8000) -> ExtractedContent:
        """Try to extract as text, otherwise return generic info."""
        # Try to read as text first
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                sample = f.read(1024)
                # Check if it looks like text
                if all(c.isprintable() or c.isspace() for c in sample[:100]):
                    # It's text - delegate to TextExtractor
                    return TextExtractor().extract(file_path, max_pages, max_text_length)
        except (UnicodeDecodeError, PermissionError, OSError):
            pass
        
        # Return generic info
        stat = file_path.stat()
        metadata = {
            "size": stat.st_size,
            "size_human": self._format_size(stat.st_size)
        }
        
        return ExtractedContent(
            file_path=file_path,
            doc_type=DocumentType.UNKNOWN,
            text_content="",
            metadata=metadata,
            extraction_method="none",
            error=f"Unknown file type: {file_path.suffix}"
        )


class ContentExtractor:
    """
    Main content extractor using Strategy Pattern.
    
    Each document type has its own extractor class that can be
    registered or replaced without modifying this class.
    
    Usage:
        extractor = ContentExtractor()
        # Add custom extractor
        extractor.register_extractor(MyCustomExtractor())
        # Extract content
        result = extractor.extract("document.pdf")
    """
    
    def __init__(self, extractors: Optional[List[DocumentExtractor]] = None):
        """
        Initialize with extractors.
        
        Args:
            extractors: List of extractors. If None, uses default set.
        """
        self.extractors: List[DocumentExtractor] = extractors or self._default_extractors()
        self.max_pages = 5
        self.max_text_length = 8000
        logger.info(f"ContentExtractor initialized with {len(self.extractors)} extractors")
    
    def _default_extractors(self) -> List[DocumentExtractor]:
        """Get the default set of extractors."""
        return [
            PDFExtractor(),
            ImageExtractor(),
            TextExtractor(),
            OfficeExtractor(),
            GenericExtractor(),  # Fallback
        ]
    
    def register_extractor(self, extractor: DocumentExtractor, position: Optional[int] = None):
        """
        Register a new extractor.
        
        Args:
            extractor: The extractor to add
            position: Optional position in list (None = append)
        """
        if position is None:
            self.extractors.append(extractor)
        else:
            self.extractors.insert(position, extractor)
        logger.info(f"Registered extractor: {type(extractor).__name__}")
    
    def unregister_extractor(self, extractor_type: type):
        """
        Remove an extractor by type.
        
        Args:
            extractor_type: The class type to remove
        """
        self.extractors = [e for e in self.extractors if not isinstance(e, extractor_type)]
        logger.info(f"Unregistered extractor: {extractor_type.__name__}")
    
    def extract(
        self,
        file_path: Union[str, Path],
        base_path: Optional[Path] = None
    ) -> ExtractedContent:
        """
        Extract content from a document.
        
        Args:
            file_path: Path to document
            base_path: Optional base directory to enforce path containment
            
        Returns:
            ExtractedContent with text and metadata
        """
        file_path = Path(file_path).resolve()
        
        # Validate path is within base directory if specified
        if base_path:
            base_path = Path(base_path).resolve()
            try:
                file_path.relative_to(base_path)
            except ValueError:
                return ExtractedContent(
                    file_path=file_path,
                    doc_type=DocumentType.UNKNOWN,
                    text_content="",
                    metadata={},
                    extraction_method="none",
                    error="Access denied: file outside allowed directory"
                )
        
        # Check for symlink attacks
        if file_path.is_symlink():
            return ExtractedContent(
                file_path=file_path,
                doc_type=DocumentType.UNKNOWN,
                text_content="",
                metadata={},
                extraction_method="none",
                error="Access denied: symlinks not allowed"
            )
        
        if not file_path.exists():
            return ExtractedContent(
                file_path=file_path,
                doc_type=DocumentType.UNKNOWN,
                text_content="",
                metadata={},
                extraction_method="none",
                error=f"File not found: {file_path}"
            )
        
        # Find appropriate extractor
        for extractor in self.extractors:
            try:
                if extractor.can_extract(file_path):
                    return extractor.extract(file_path, self.max_pages, self.max_text_length)
            except Exception as e:
                logger.warning(f"Extractor {type(extractor).__name__} failed: {e}")
                continue
        
        # Should never reach here due to GenericExtractor, but just in case
        return ExtractedContent(
            file_path=file_path,
            doc_type=DocumentType.UNKNOWN,
            text_content="",
            metadata={},
            extraction_method="none",
            error="No suitable extractor found"
        )
    
    def get_supported_types(self) -> Dict[str, List[str]]:
        """Get list of supported file types by extractor."""
        types = {}
        for extractor in self.extractors:
            name = type(extractor).__name__
            exts = getattr(extractor, 'SUPPORTED_EXTENSIONS', set())
            types[name] = list(exts)
        return types
    
    def get_installation_instructions(self) -> Dict[str, str]:
        """Get installation instructions for missing dependencies."""
        return {
            "tesseract": "brew install tesseract (macOS) or apt-get install tesseract-ocr (Linux)",
            "pymupdf": "pip install pymupdf",
            "pdfplumber": "pip install pdfplumber",
            "pdf2image": "pip install pdf2image (also requires poppler: brew install poppler)",
            "pytesseract": "pip install pytesseract",
            "python-docx": "pip install python-docx",
            "pillow": "pip install Pillow"
        }
